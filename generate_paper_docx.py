#!/usr/bin/env python3
"""
generate_paper_docx.py
======================
Generates a publication-quality DOCX manuscript for:
"Disentangling Natural and Anthropogenic Drivers of Global Terrestrial
 Water Storage Decline using Machine Learning"

Figures are NOT embedded — instead, clearly marked placeholders describe
each figure so the author can insert them manually.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# 0. CONFIGURATION
# ---------------------------------------------------------------------------
OUTPUT_DIR = r"c:\SILIKA\Thesis\data2"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "TWS_Attribution_Manuscript.docx")

doc = Document()

# ---------------------------------------------------------------------------
# STYLE CONFIGURATION
# ---------------------------------------------------------------------------
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0  # Double spacing
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Heading styles
for level in range(1, 4):
    h_style = doc.styles[f"Heading {level}"]
    h_font = h_style.font
    h_font.name = "Times New Roman"
    h_font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        h_font.size = Pt(14)
        h_font.bold = True
    elif level == 2:
        h_font.size = Pt(13)
        h_font.bold = True
    else:
        h_font.size = Pt(12)
        h_font.bold = True
        h_font.italic = True


def add_paragraph(text, bold=False, italic=False, alignment=None, space_after=6):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _set_cell_borders_and_shading(cell, border_color, fill_color):
    """Set cell borders and background shading using proper OxmlElement."""
    from lxml import etree
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Build tcBorders element
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    w = nsmap["w"]
    tcBorders = etree.SubElement(tcPr, qn("w:tcBorders"))
    for side in ["top", "left", "bottom", "right"]:
        etree.SubElement(tcBorders, qn(f"w:{side}"), attrib={
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): border_color,
        })
    
    # Shading
    etree.SubElement(tcPr, qn("w:shd"), attrib={
        qn("w:fill"): fill_color,
        qn("w:val"): "clear",
    })


def add_figure_placeholder(fig_number, caption, description):
    """Add a clearly marked figure placeholder box."""
    # Add a bordered paragraph as a visual placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    # Create border effect using a table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    _set_cell_borders_and_shading(cell, "808080", "F0F0F0")

    # Add placeholder text
    placeholder_p = cell.paragraphs[0]
    placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = placeholder_p.add_run(f"\n[INSERT {fig_number} HERE]\n")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(128, 128, 128)

    desc_p = cell.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc_p.add_run(f"Description: {description}")
    desc_run.italic = True
    desc_run.font.size = Pt(10)
    desc_run.font.name = "Times New Roman"
    desc_run.font.color.rgb = RGBColor(100, 100, 100)

    cell.add_paragraph()  # spacing

    # Caption below the placeholder
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run(f"{fig_number}. ")
    cap_run.bold = True
    cap_run.font.size = Pt(11)
    cap_run.font.name = "Times New Roman"
    cap_run2 = cap_p.add_run(caption)
    cap_run2.font.size = Pt(11)
    cap_run2.font.name = "Times New Roman"


def add_table_placeholder(table_number, caption, description):
    """Add a clearly marked table placeholder box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    _set_cell_borders_and_shading(cell, "4472C4", "E8F0FE")

    placeholder_p = cell.paragraphs[0]
    placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = placeholder_p.add_run(f"\n[INSERT {table_number} HERE]\n")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(68, 114, 196)

    desc_p = cell.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc_p.add_run(f"Description: {description}")
    desc_run.italic = True
    desc_run.font.size = Pt(10)
    desc_run.font.name = "Times New Roman"
    desc_run.font.color.rgb = RGBColor(80, 80, 120)

    cell.add_paragraph()

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run(f"{table_number}. ")
    cap_run.bold = True
    cap_run.font.size = Pt(11)
    cap_run.font.name = "Times New Roman"
    cap_run2 = cap_p.add_run(caption)
    cap_run2.font.size = Pt(11)
    cap_run2.font.name = "Times New Roman"


def add_equation(eq_text, eq_number):
    """Add a numbered equation."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(eq_text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Add equation number right-aligned using tab
    tab_run = p.add_run(f"\t\t\t({eq_number})")
    tab_run.font.name = "Times New Roman"
    tab_run.font.size = Pt(12)
    return p


# ===========================================================================
# TITLE PAGE
# ===========================================================================
doc.add_paragraph()  # spacing

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(72)
title_p.paragraph_format.space_after = Pt(24)
title_run = title_p.add_run(
    "Disentangling Natural and Anthropogenic Drivers of "
    "Global Terrestrial Water Storage Decline using Machine Learning"
)
title_run.bold = True
title_run.font.size = Pt(16)
title_run.font.name = "Times New Roman"

# Author placeholder
auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.paragraph_format.space_after = Pt(6)
auth_run = auth_p.add_run("[Author Name(s)]")
auth_run.font.size = Pt(13)
auth_run.font.name = "Times New Roman"

# Affiliation placeholder
aff_p = doc.add_paragraph()
aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff_p.paragraph_format.space_after = Pt(6)
aff_run = aff_p.add_run("[Affiliation(s), Department, University/Institution]")
aff_run.italic = True
aff_run.font.size = Pt(11)
aff_run.font.name = "Times New Roman"

# Corresponding author
corr_p = doc.add_paragraph()
corr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
corr_p.paragraph_format.space_after = Pt(36)
corr_run = corr_p.add_run("Corresponding Author: [email@institution.edu]")
corr_run.font.size = Pt(11)
corr_run.font.name = "Times New Roman"

doc.add_page_break()

# ===========================================================================
# ABSTRACT
# ===========================================================================
doc.add_heading("Abstract", level=1)

add_paragraph(
    "Terrestrial Water Storage (TWS), encompassing surface water, soil moisture, "
    "groundwater, and snow/ice reserves, is declining at unprecedented rates across "
    "numerous global river basins, posing severe threats to water security, food "
    "production, and ecosystem sustainability. While satellite gravimetry missions\u2014"
    "specifically the Gravity Recovery and Climate Experiment (GRACE, 2002\u20132017) and "
    "its successor GRACE Follow-On (GRACE-FO, 2018\u2013present)\u2014have revolutionized our "
    "ability to monitor large-scale water mass redistribution, disentangling the "
    "relative contributions of natural hydroclimate variability from anthropogenic "
    "water abstractions remains a formidable challenge. Here, we present a Twin "
    "Machine Learning Attribution Framework based on Random Forest ensemble regressors, "
    "applied across the world\u2019s 103 largest river basins at 0.5\u00b0 \u00d7 0.5\u00b0 spatial "
    "resolution. We first reconstruct the ~11-month observational gap between GRACE "
    "and GRACE-FO (July 2017\u2013May 2018) using hydroclimate predictor variables "
    "(precipitation, evapotranspiration, and runoff) from ERA5 reanalysis and GLEAM "
    "datasets. We then quantify long-term TWS trends using the robust Theil-Sen slope "
    "estimator and assess their statistical significance via the autocorrelation-corrected "
    "Hamed and Rao Modified Mann-Kendall test. Our twin attribution models\u2014a Natural "
    "Baseline Model (M_nat) driven solely by climate fluxes and a Full Anthropogenic "
    "Model (M_anthro) augmented with groundwater and surface water abstraction rates "
    "from PCR-GLOBWB\u2014enable quantification of the Variance Explained Gain (\u0394R\u00b2) "
    "attributable to human water abstractions. Results reveal that [XX]% of the 103 "
    "basins exhibit statistically significant TWS decline trends. In heavily stressed "
    "basins such as the Indus, Tigris-Euphrates, and [other basins], groundwater "
    "abstraction emerges as the dominant driver, with \u0394R\u00b2 exceeding [XX]. These "
    "findings provide actionable, basin-specific attribution diagnostics to inform "
    "targeted water resource management and climate adaptation strategies."
)

# Keywords
kw_p = doc.add_paragraph()
kw_run1 = kw_p.add_run("Keywords: ")
kw_run1.bold = True
kw_run1.font.name = "Times New Roman"
kw_run1.font.size = Pt(12)
kw_run2 = kw_p.add_run(
    "Terrestrial Water Storage; GRACE/GRACE-FO; Machine Learning; "
    "Random Forest; Driver Attribution; Groundwater Depletion; "
    "Theil-Sen Slope; Modified Mann-Kendall Test"
)
kw_run2.font.name = "Times New Roman"
kw_run2.font.size = Pt(12)

doc.add_page_break()

# ===========================================================================
# 1. INTRODUCTION
# ===========================================================================
doc.add_heading("1. Introduction", level=1)

add_paragraph(
    "Freshwater availability is among the most critical determinants of human "
    "well-being, agricultural productivity, and ecosystem health. Terrestrial Water "
    "Storage (TWS)\u2014the vertically integrated sum of all water stored on and beneath "
    "the Earth\u2019s land surface, including surface water reservoirs, soil moisture, "
    "groundwater aquifers, and snow/ice\u2014represents the ultimate buffer against "
    "hydroclimatic extremes. Over the past two decades, multiple lines of evidence "
    "have converged to reveal an alarming global pattern: TWS is declining in many "
    "of the world\u2019s most productive and densely populated river basins (Rodell et al., "
    "2018; Scanlon et al., 2018; Famiglietti, 2014)."
)

add_paragraph(
    "The Gravity Recovery and Climate Experiment (GRACE) mission, launched in March "
    "2002, and its successor GRACE Follow-On (GRACE-FO), operational since June 2018, "
    "have provided an unprecedented observational record of monthly variations in Earth\u2019s "
    "gravity field, from which changes in TWS can be inferred at spatial scales of "
    "~300 km and larger (Tapley et al., 2004; Landerer and Swenson, 2012). These "
    "satellite gravimetry observations have unequivocally demonstrated massive water "
    "depletions in regions including the Indo-Gangetic Plain (Tiwari et al., 2009), "
    "the Central Valley of California (Famiglietti et al., 2011), the Middle East "
    "(Voss et al., 2013), and the North China Plain (Feng et al., 2013). The GRACE "
    "record, however, contains an approximately 11-month observational gap between "
    "the end of the original GRACE mission (June 2017) and the start of GRACE-FO "
    "(June 2018), which must be bridged for continuous trend analysis."
)

add_paragraph(
    "A central and unresolved challenge in the interpretation of TWS variability lies "
    "in the attribution of observed changes to their underlying drivers. TWS dynamics "
    "are governed by the terrestrial water balance:"
)

add_equation(
    "dTWS/dt  =  TWSC  =  P  \u2212  ET  \u2212  Q  \u2212  (GW_abs + SW_abs)", 1
)

add_paragraph(
    "where TWSC denotes the Terrestrial Water Storage Change, P is precipitation, "
    "ET is evapotranspiration, Q is runoff/discharge, and GW_abs and SW_abs represent "
    "groundwater and surface water abstraction rates, respectively. Observed TWS declines "
    "may thus reflect long-term precipitation deficits (e.g., megadroughts exacerbated "
    "by climate change), enhanced evapotranspiration driven by rising temperatures, "
    "or unsustainable human water withdrawals for irrigation, industrial use, and "
    "municipal supply. In practice, these drivers often operate simultaneously and "
    "interact nonlinearly, making their separation particularly challenging when relying "
    "on simple correlation or regression approaches."
)

add_paragraph(
    "Previous studies have employed hydrological models (e.g., D\u00f6ll et al., 2014), "
    "data assimilation frameworks (e.g., Zaitchik et al., 2008), and statistical "
    "decomposition methods (e.g., Humphrey et al., 2017) to attribute TWS variability. "
    "However, most approaches suffer from one or more limitations: reliance on a single "
    "hydrological model with inherent structural uncertainties, inability to separate "
    "climate and anthropogenic signals at the basin scale, or lack of a rigorous "
    "statistical framework for trend significance testing that accounts for temporal "
    "autocorrelation in hydrological time series."
)

add_paragraph(
    "In this study, we address these gaps through a comprehensive, data-driven "
    "attribution framework comprising three methodological pillars:"
)

add_paragraph(
    "(i) Machine learning gap-filling: We reconstruct the GRACE\u2013GRACE-FO observational "
    "gap using Random Forest regressors trained on coincident hydroclimate predictor "
    "variables, yielding a continuous TWS record spanning April 2002 to December 2019 "
    "across 103 major global river basins.",
    space_after=3,
)

add_paragraph(
    "(ii) Robust trend estimation: We quantify long-term TWS trends using the "
    "non-parametric Theil-Sen slope estimator and assess their statistical significance "
    "using the Hamed and Rao (1998) Modified Mann-Kendall test, which explicitly "
    "corrects for temporal autocorrelation\u2014a critical but often neglected consideration "
    "in hydrological trend analysis.",
    space_after=3,
)

add_paragraph(
    "(iii) Twin machine learning attribution: We deploy paired Random Forest models\u2014a "
    "Natural Baseline Model (M_nat) driven exclusively by climate variables (P, ET, Q) "
    "and a Full Anthropogenic Model (M_anthro) augmented with groundwater and surface "
    "water abstraction rates\u2014to isolate the explanatory power gained by including human "
    "water use variables, quantified through the Variance Explained Gain (\u0394R\u00b2) and "
    "Out-of-Bag permutation feature importance."
)

doc.add_page_break()

# ===========================================================================
# 2. DATA AND METHODS
# ===========================================================================
doc.add_heading("2. Data and Methods", level=1)

# --- 2.1 Study Domain ---
doc.add_heading("2.1. Study Domain", level=2)

add_paragraph(
    "The analysis encompasses the world\u2019s 103 largest river basins, which collectively "
    "drain approximately 75% of the global continental land area and support the majority "
    "of the world\u2019s population. Basin boundaries are delineated using a high-resolution "
    "basin mask at 0.5\u00b0 \u00d7 0.5\u00b0 spatial resolution (720 \u00d7 360 grid cells), consistent "
    "with the native resolution of the input datasets. All spatial aggregations employ "
    "latitude-cosine area weighting (cos(lat)) to account for the decrease in grid cell "
    "area toward the poles, ensuring physically meaningful basin-average estimates."
)

add_figure_placeholder(
    "Figure 1",
    "Study domain showing the 103 largest global river basins used in this analysis. "
    "Basins are shaded by unique identifiers overlaid on a global land mask at "
    "0.5\u00b0 \u00d7 0.5\u00b0 resolution. Major basins discussed in the text (Indus, "
    "Tigris-Euphrates, Colorado, Amazon, Congo) are labeled.",
    "A global map (Robinson or Mollweide projection) displaying all 103 river basin "
    "polygons, color-coded by basin ID or grouped by continent. Include latitude/longitude "
    "grid lines, a scale bar, and labels for the key basins discussed in the paper.",
)

# --- 2.2 Datasets ---
doc.add_heading("2.2. Datasets", level=2)

doc.add_heading("2.2.1. GRACE/GRACE-FO Terrestrial Water Storage Anomalies", level=3)

add_paragraph(
    "Monthly TWS anomaly fields are derived from the ensemble mean of three independent "
    "GRACE/GRACE-FO processing centers: the Jet Propulsion Laboratory (JPL), GeoForschungsZentrum "
    "(GFZ), and the Center for Space Research (CSR). TWS anomalies are expressed as Liquid "
    "Water Equivalent (LWE) thickness in centimeters relative to a 2004\u20132009 baseline "
    "period. The native GRACE temporal coverage extends from April 2002 through June 2017, "
    "with GRACE-FO resuming in June 2018. Missing months within the operational periods "
    "(instrument anomalies, orbit maneuvers) and the ~11-month inter-mission gap (July "
    "2017\u2013May 2018) are explicitly represented as NaN values in the analysis timeline, "
    "pending reconstruction (Section 2.4)."
)

doc.add_heading("2.2.2. ERA5 Reanalysis: Precipitation and Runoff", level=3)

add_paragraph(
    "Monthly total precipitation (P) and total runoff (Q) fields are obtained from the "
    "European Centre for Medium-Range Weather Forecasts (ECMWF) ERA5 reanalysis product "
    "(Hersbach et al., 2020). ERA5 provides globally complete, hourly atmospheric fields "
    "at 0.25\u00b0 resolution, aggregated here to monthly totals on a 0.5\u00b0 \u00d7 0.5\u00b0 grid. "
    "Precipitation is converted from meters/month to centimeters/month (\u00d7100), and runoff "
    "is similarly converted from meters/month to centimeters/month to maintain dimensional "
    "consistency across all water balance terms."
)

doc.add_heading("2.2.3. GLEAM Evapotranspiration", level=3)

add_paragraph(
    "Actual evapotranspiration (ET) estimates are sourced from the Global Land Evaporation "
    "Amsterdam Model (GLEAM) v3.x (Martens et al., 2017; Miralles et al., 2011). GLEAM "
    "provides satellite-observation-driven ET estimates at 0.25\u00b0 resolution, regridded to "
    "0.5\u00b0 \u00d7 0.5\u00b0 for consistency. The native GLEAM units of millimeters/month are converted "
    "to centimeters/month (\u00f710). Spatial alignment requires a latitude flip (North\u2192South to "
    "South\u2192North) and a 180\u00b0 longitudinal shift to match the \u2212180\u00b0 to +180\u00b0 convention "
    "used throughout this analysis."
)

doc.add_heading("2.2.4. PCR-GLOBWB Groundwater and Surface Water Abstraction", level=3)

add_paragraph(
    "Spatially explicit estimates of groundwater abstraction (GW_abs) and surface water "
    "abstraction (SW_abs) are obtained from the PCR-GLOBWB 2.0 global hydrological model "
    "(Sutanudjaja et al., 2018). PCR-GLOBWB simulates global water demand, allocation, and "
    "abstraction from both groundwater and surface water sources at 0.5\u00b0 resolution, "
    "accounting for irrigation, domestic, and industrial water use sectors. Abstraction "
    "rates are converted from meters/month to centimeters/month (\u00d7100)."
)

add_table_placeholder(
    "Table 1",
    "Summary of input datasets used in this study, including variable names, native units, "
    "converted units (cm/month), spatial resolution, temporal coverage, and source references.",
    "A multi-column table with rows for each dataset: GRACE/GRACE-FO TWS, ERA5 P, ERA5 Q, "
    "GLEAM ET, PCR-GLOBWB GW_abs, PCR-GLOBWB SW_abs. Columns: Variable, Source, Native "
    "Resolution, Unit Conversion, Temporal Coverage, Reference.",
)

# --- 2.3 Data Preprocessing ---
doc.add_heading("2.3. Data Preprocessing and Unit Harmonization", level=2)

add_paragraph(
    "All gridded datasets are preprocessed to ensure spatial, temporal, and dimensional "
    "consistency prior to analysis. The preprocessing pipeline (implemented in MATLAB "
    "R2021b) proceeds as follows:"
)

add_paragraph(
    "Non-physical fill values: All NetCDF variables are scanned for non-physical fill "
    "values (\u2264 \u2212900, > 1 \u00d7 10\u00b9\u2079, \u00b1Inf) and replaced with MATLAB NaN immediately upon "
    "ingestion, preventing contamination of downstream calculations."
)

add_paragraph(
    "Unit harmonization: All hydroclimate fluxes (P, ET, Q, GW_abs, SW_abs) and TWS "
    "anomalies are converted to a common unit of liquid water height equivalent depth "
    "(cm/month) prior to any spatial aggregation or modeling."
)

add_paragraph(
    "Temporal alignment: All datasets are re-indexed onto a common 213-month timeline "
    "(April 2002 through December 2019). GRACE months that are observationally present are "
    "mapped using their native timestamps; missing months are encoded as NaN for subsequent "
    "gap-filling."
)

add_paragraph(
    "Area-weighted basin aggregation: Gridded fields are spatially aggregated over each "
    "of the 103 river basins using latitude-cosine weighting. For each basin b and variable "
    "X, the basin-average time series is computed as:"
)

add_equation(
    "X\u0305_b(t)  =  \u03a3\u1d62 [ X(i, t) \u00d7 cos(lat\u1d62) ] / \u03a3\u1d62 [ cos(lat\u1d62) ]      "
    "for all grid cells i \u2208 basin b",
    2,
)

add_paragraph(
    "where lat\u1d62 is the latitude of grid cell i. This yields six basin-average time series "
    "matrices of dimension [213 \u00d7 103]: TWS, P, ET, Q, GW_abs, and SW_abs."
)

# --- 2.4 Gap-Filling ---
doc.add_heading("2.4. Machine Learning Gap-Filling of GRACE Observations", level=2)

add_paragraph(
    "To obtain a continuous TWS record, we reconstruct missing GRACE months\u2014including "
    "the ~11-month inter-mission gap (July 2017\u2013May 2018) and sporadic within-mission "
    "dropouts\u2014using Random Forest (RF) ensemble regressors (Breiman, 2001). For each "
    "basin b, a separate RF model is trained on the subset of months where GRACE TWS "
    "observations are available and predictor variables (P, ET, Q, and the water balance "
    "residual P \u2212 ET \u2212 Q) are non-missing. The trained model is then applied to predict "
    "TWS anomalies at all gap months where predictors are available."
)

add_paragraph(
    "RF hyperparameters are set to 200 trees (n_trees = 200) with a minimum leaf size "
    "of 5 (min_leaf_size = 5) to balance predictive accuracy and regularization against "
    "overfitting. Model skill is evaluated using Out-of-Bag (OOB) predictions, which "
    "provide an unbiased estimate of generalization error without requiring a separate "
    "validation set. The OOB Root Mean Square Error (RMSE) and coefficient of determination "
    "(R\u00b2) are recorded for each basin. All 103 basins are processed in parallel using "
    "MATLAB\u2019s parfor construct for computational efficiency on the DIRAC HPC cluster."
)

add_figure_placeholder(
    "Figure 2",
    "Validation of the Random Forest gap-filling model. (a) Scatter plot of OOB-predicted "
    "versus observed TWS anomalies for a representative basin (e.g., Indus or Amazon), "
    "with 1:1 reference line, R\u00b2, and RMSE annotated. (b) Time series of observed GRACE "
    "TWS (circles), gap-filled TWS (red line), and the shaded inter-mission gap period "
    "(July 2017\u2013May 2018) for the same basin. (c) Histogram or boxplot of OOB R\u00b2 values "
    "across all 103 basins.",
    "Three-panel figure. Panel (a): scatter plot with color-coded density. Panel (b): "
    "time series plot showing observed TWS with markers, continuous reconstructed TWS as a "
    "line, and a gray or red shaded region for the 2017\u20132018 gap. Panel (c): distribution "
    "of OOB R\u00b2 across all 103 basins showing the quality of reconstruction.",
)

# --- 2.5 TWSC Calculation ---
doc.add_heading("2.5. Computation of TWS Change (TWSC)", level=2)

add_paragraph(
    "The rate of change of TWS (TWSC) is computed from the gap-filled, continuous TWS "
    "time series using centered finite differences:"
)

add_equation(
    "TWSC(t)  =  [ TWS(t+1) \u2212 TWS(t\u22121) ] / (2\u0394t)", 3
)

add_paragraph(
    "where \u0394t = 1 month. Forward and backward differences are applied at the first and "
    "last time steps, respectively. Prior to attribution modeling, all variables (TWSC, P, "
    "ET, Q, GW_abs, SW_abs) are deseasonalized by subtracting the long-term monthly "
    "climatological mean, removing the dominant seasonal cycle and isolating interannual "
    "to decadal anomaly signals."
)

# --- 2.6 Twin Attribution Framework ---
doc.add_heading("2.6. Twin Random Forest Attribution Framework", level=2)

add_paragraph(
    "The core methodological innovation of this study is a twin machine learning framework "
    "designed to isolate the explanatory contribution of anthropogenic water abstractions "
    "to observed TWS variability. Two Random Forest regression models are trained "
    "independently for each basin:"
)

add_paragraph(
    "Model 1\u2014Natural Baseline (M_nat): Predicts TWSC as a function of natural "
    "hydroclimate drivers alone:",
    bold=True,
    space_after=3,
)
add_equation("TWSC_pred  =  f_nat(P, ET, Q)", 4)

add_paragraph(
    "Model 2\u2014Full Anthropogenic (M_anthro): Predicts TWSC as a function of both "
    "natural and anthropogenic drivers:",
    bold=True,
    space_after=3,
)
add_equation("TWSC_pred  =  f_anthro(P, ET, Q, GW_abs, SW_abs)", 5)

add_paragraph(
    "Both models use 200 trees with a minimum leaf size of 10 and are evaluated using "
    "OOB predictions. The OOB R\u00b2 is computed for each model, and the Variance Explained "
    "Gain is defined as:"
)

add_equation("\u0394R\u00b2  =  R\u00b2_anthro  \u2212  R\u00b2_nat", 6)

add_paragraph(
    "A positive \u0394R\u00b2 indicates that human water abstractions explain a statistically "
    "meaningful fraction of TWS variability beyond what can be attributed to natural "
    "climate forcing. To identify the single most important driver for each basin, "
    "we extract Out-of-Bag Permutation Feature Importance scores from M_anthro, which "
    "quantify the increase in OOB prediction error when each feature\u2019s values are "
    "randomly permuted, thereby breaking its relationship with the target variable."
)

add_figure_placeholder(
    "Figure 3",
    "Schematic diagram of the Twin Random Forest Attribution Framework. Left branch: "
    "M_nat (Natural Baseline) trained on P, ET, Q. Right branch: M_anthro (Full "
    "Anthropogenic) trained on P, ET, Q, GW_abs, SW_abs. The Variance Explained Gain "
    "(\u0394R\u00b2) is computed as the difference in OOB R\u00b2 between the two models.",
    "A conceptual/methodological flowchart showing the twin model architecture. Input "
    "boxes for each predictor variable, two parallel Random Forest model boxes, output "
    "arrows to TWSC predictions, and the \u0394R\u00b2 comparison step. Use a clean, schematic "
    "style suitable for publication.",
)

# --- 2.7 Trend Estimation ---
doc.add_heading("2.7. Trend Estimation and Significance Testing", level=2)

add_paragraph(
    "Long-term TWS trends are estimated using the Theil-Sen slope estimator (Theil, 1950; "
    "Sen, 1968), a non-parametric method that computes the median of all pairwise slopes "
    "between data points, providing robustness against outliers and non-normality. The "
    "estimated slope \u03b2\u0302 (in cm/year) represents the rate of TWS change over the analysis "
    "period."
)

add_paragraph(
    "The statistical significance of each basin\u2019s trend is assessed using the Modified "
    "Mann-Kendall (MK) test of Hamed and Rao (1998), which corrects the variance of the "
    "Mann-Kendall S statistic for temporal autocorrelation in the data. The correction "
    "factor accounts for significant autocorrelation coefficients (at the 95% confidence "
    "level) up to lag n/4, where n is the series length. A trend is declared statistically "
    "significant at the \u03b1 = 0.05 level. This autocorrelation correction is critical for "
    "hydrological time series, where persistence (e.g., multi-year droughts, sustained "
    "pumping) can inflate apparent trend significance if standard MK tests are na\u00efvely "
    "applied."
)

# --- 2.8 Block CV ---
doc.add_heading("2.8. Model Validation: 3-Year Contiguous Block Cross-Validation", level=2)

add_paragraph(
    "To rigorously evaluate the predictive skill of both M_nat and M_anthro without "
    "temporal information leakage, we employ 3-Year Contiguous Block Cross-Validation "
    "(Block CV). Standard random K-fold cross-validation is inappropriate for "
    "autocorrelated time series because random partitioning allows temporally adjacent "
    "observations to appear in both training and test sets, artificially inflating "
    "performance metrics. In Block CV, the 213-month record is divided into non-overlapping "
    "contiguous blocks of 36 months (3 years). For each fold, one block is held out as the "
    "test set while all remaining blocks serve as the training set. Models are trained de "
    "novo on the training blocks and evaluated on the held-out test block."
)

add_paragraph(
    "Four hydrologic performance metrics are computed for each basin across all Block CV "
    "folds: (i) Nash-Sutcliffe Efficiency (NSE), (ii) Kling-Gupta Efficiency (KGE, using "
    "a modified formulation suitable for zero-mean anomaly variables), (iii) Root Mean "
    "Square Error (RMSE, cm/month), and (iv) Pearson R\u00b2. These metrics collectively assess "
    "correlation, bias, variability ratio, and absolute error magnitude."
)

doc.add_page_break()

# ===========================================================================
# 3. RESULTS
# ===========================================================================
doc.add_heading("3. Results", level=1)

# --- 3.1 Global TWS Trends ---
doc.add_heading("3.1. Global Patterns of TWS Decline", level=2)

add_paragraph(
    "Application of the Theil-Sen slope estimator to the gap-filled, continuous TWS record "
    "(April 2002\u2013December 2019) reveals widespread negative TWS trends across the 103 "
    "basins analyzed (Figure 4). [XX] out of 103 basins ([XX]%) exhibit statistically "
    "significant (p < 0.05, Modified Mann-Kendall test) negative trends, with a mean "
    "decline rate of [XX] cm/year among these basins. The most severe depletions are "
    "observed in [basin names], with trend magnitudes exceeding [XX] cm/year. Conversely, "
    "[XX] basins show significant positive TWS trends, primarily located in [regions], "
    "likely reflecting increased precipitation or glacier/snowmelt contributions."
)

add_figure_placeholder(
    "Figure 4",
    "Global map of TWS trends (cm/year) across the 103 largest river basins for the period "
    "April 2002\u2013December 2019. A diverging blue (positive/gaining)\u2013red (negative/declining) "
    "color scale centered at zero is used. Basins with statistically non-significant trends "
    "(p \u2265 0.05, Hamed\u2013Rao Modified Mann-Kendall test) are overlaid with gray stippling. "
    "Non-basin areas are shown in light gray.",
    "A global choropleth map with each basin polygon colored by its Theil-Sen TWS trend slope. "
    "Use a diverging Red-White-Blue colormap symmetric about zero. Overlay stipple dots on "
    "non-significant basins. Include a colorbar labeled 'TWS Trend (cm/year)' and a title. "
    "Generated by plot_basin_trends.m.",
)

add_figure_placeholder(
    "Figure 5",
    "Bar chart of the 20 basins with the most negative TWS trends (cm/year), ranked from "
    "most severe to least severe. Bars are colored red for statistically significant trends "
    "and gray for non-significant trends. Error bars or confidence intervals from the "
    "Theil-Sen estimator may be included.",
    "A horizontal or vertical bar chart showing the top-20 declining basins with their "
    "Theil-Sen trend slopes. Basin names or IDs on the y-axis. Significance indicated by "
    "bar color.",
)

# --- 3.2 Model Validation ---
doc.add_heading("3.2. Validation of Machine Learning Models", level=2)

add_paragraph(
    "The 3-Year Contiguous Block Cross-Validation demonstrates that both twin attribution "
    "models achieve robust predictive skill across the majority of basins. The Natural "
    "Baseline Model (M_nat) achieves a median Block CV NSE of [XX] (interquartile range: "
    "[XX\u2013XX]), median KGE of [XX], and median RMSE of [XX] cm/month. The Full Anthropogenic "
    "Model (M_anthro) achieves improved performance with a median NSE of [XX], median KGE "
    "of [XX], and median RMSE of [XX] cm/month (Table 2). These results confirm that the "
    "models capture the dominant hydrological dynamics without overfitting, as the Block CV "
    "procedure explicitly prevents temporal leakage."
)

add_table_placeholder(
    "Table 2",
    "Summary of 3-Year Block Cross-Validation performance metrics for the Natural Baseline "
    "(M_nat) and Full Anthropogenic (M_anthro) models across all 103 basins. Metrics include "
    "median, mean, interquartile range, and percentage of basins with NSE > 0 (skillful).",
    "A table with rows: NSE, KGE, RMSE (cm/month), R\u00b2. Columns grouped by M_nat and M_anthro, "
    "each showing Median, Mean, IQR, and % basins with NSE > 0.",
)

add_figure_placeholder(
    "Figure 6",
    "Box-and-whisker plots comparing Block Cross-Validation performance metrics (NSE, KGE, "
    "RMSE) between the Natural Baseline (M_nat, blue) and Full Anthropogenic (M_anthro, "
    "orange) models across all 103 basins. Whiskers extend to the 5th and 95th percentiles; "
    "outliers are plotted as individual points.",
    "Side-by-side boxplots for each metric. Two boxes per metric (M_nat and M_anthro). "
    "NSE and KGE should show that M_anthro generally matches or exceeds M_nat. RMSE should "
    "show that M_anthro generally has equal or lower error.",
)

# --- 3.3 Attribution ---
doc.add_heading("3.3. Attribution: Disentangling Natural and Anthropogenic Drivers", level=2)

add_paragraph(
    "The twin attribution framework reveals heterogeneous driver dominance across the "
    "global basin network. Figure 7 presents a global map of the dominant driver for each "
    "basin, defined as the feature with the highest OOB permutation importance in M_anthro. "
    "Precipitation (P) emerges as the dominant driver in [XX]% of basins, predominantly in "
    "tropical and high-latitude regions. Evapotranspiration (ET) dominates in [XX]% of "
    "basins, particularly in semi-arid and warming regions. Groundwater abstraction (GW_abs) "
    "is identified as the dominant driver in [XX]% of basins, concentrated in South Asia "
    "(Indus, Ganges), the Middle East (Tigris-Euphrates), and parts of North America and "
    "North Africa."
)

add_figure_placeholder(
    "Figure 7",
    "Global map of the dominant driver of TWS variability in each of the 103 basins, "
    "determined by maximum OOB permutation feature importance from M_anthro. Basins are "
    "color-coded by dominant driver category: Precipitation (blue), Evapotranspiration "
    "(green), Runoff (cyan), Groundwater Abstraction (orange), Surface Water Abstraction "
    "(red). A categorical legend is provided.",
    "A global choropleth map with each basin polygon colored by its dominant driver "
    "(5 categories). Use the colormap: P=Blue, ET=Green, Q=Cyan, GW=Orange, SW=Red. "
    "Include a categorical colorbar/legend. Generated by plot_global_attribution_map.m.",
)

add_paragraph(
    "The Variance Explained Gain (\u0394R\u00b2) provides a quantitative measure of the additional "
    "explanatory power contributed by anthropogenic abstractions. Across all 103 basins, "
    "the mean \u0394R\u00b2 is [XX] (\u00b1 [XX] standard deviation). However, the distribution is highly "
    "skewed: basins with known intensive groundwater exploitation exhibit \u0394R\u00b2 values "
    "exceeding [XX], while basins with minimal human abstraction show near-zero or "
    "slightly negative \u0394R\u00b2 values (Figure 8)."
)

add_figure_placeholder(
    "Figure 8",
    "Scatter plot or bar chart of Variance Explained Gain (\u0394R\u00b2 = R\u00b2_anthro \u2212 R\u00b2_nat) for "
    "the 103 basins, sorted from highest to lowest \u0394R\u00b2. Basins are colored by their dominant "
    "driver category. Key basins (Indus, Tigris-Euphrates, Colorado, etc.) are labeled. "
    "A dashed horizontal line at \u0394R\u00b2 = 0 separates basins where anthropogenic variables "
    "improve model performance from those where they do not.",
    "A sorted bar chart or lollipop plot with basin ID/name on the x-axis and \u0394R\u00b2 on the "
    "y-axis. Bars colored by dominant driver. Label the top 10 and bottom 5 basins.",
)

add_figure_placeholder(
    "Figure 9",
    "Stacked or grouped bar chart of OOB permutation feature importance for the top 15 "
    "most stressed basins (highest negative TWS trend \u00d7 significant). Five stacked segments "
    "per bar represent P, ET, Q, GW_abs, and SW_abs importance. Basin names are shown on "
    "the y-axis.",
    "A horizontal stacked bar chart for the 15 most severely declining basins. Each bar is "
    "segmented into 5 features (P, ET, Q, GW_abs, SW_abs) with consistent colors matching "
    "Figure 7. Annotate where GW_abs or SW_abs dominate.",
)

doc.add_page_break()

# ===========================================================================
# 4. DISCUSSION
# ===========================================================================
doc.add_heading("4. Discussion", level=1)

# --- 4.1 Hotspots ---
doc.add_heading("4.1. Hotspots of Human-Induced TWS Depletion", level=2)

add_paragraph(
    "Our twin attribution framework identifies several well-documented hotspots of "
    "anthropogenic TWS depletion, lending credibility to the methodology and providing "
    "new quantitative attribution metrics."
)

add_paragraph(
    "Indus Basin: The Indus Basin, encompassing the intensively irrigated Indo-Gangetic "
    "aquifer system, exhibits one of the most severe TWS decline rates globally ([XX] "
    "cm/year, p < 0.001). The \u0394R\u00b2 for this basin is [XX], indicating that groundwater "
    "abstraction explains [XX]% of additional TWS variance beyond natural climate drivers. "
    "OOB feature importance confirms GW_abs as the dominant driver (importance = [XX]), "
    "consistent with extensive literature documenting unsustainable groundwater extraction "
    "for rice-wheat cropping systems (Tiwari et al., 2009; Rodell et al., 2009).",
    space_after=3,
)

add_paragraph(
    "Tigris-Euphrates Basin: The Tigris-Euphrates system, spanning Turkey, Syria, and Iraq, "
    "shows a significant negative TWS trend of [XX] cm/year. The attribution analysis "
    "reveals a \u0394R\u00b2 of [XX], with both GW_abs and SW_abs contributing substantially to the "
    "observed depletion. This is consistent with the combined effects of upstream dam "
    "construction, expanding irrigation networks, and drought conditions reported by Voss "
    "et al. (2013) and Joodaki et al. (2014).",
    space_after=3,
)

add_paragraph(
    "Colorado Basin: The Colorado River Basin shows [describe findings]. The \u0394R\u00b2 of [XX] "
    "suggests [interpretation]. [Discuss in relation to Castle et al. (2014) findings on "
    "Lake Mead and groundwater depletion.]",
)

add_figure_placeholder(
    "Figure 10",
    "Detailed case study panels for three key basins: (a) Indus, (b) Tigris-Euphrates, "
    "(c) Colorado. Each panel shows: (top) time series of observed and reconstructed TWS "
    "anomalies; (middle) observed TWSC versus M_nat and M_anthro predictions; (bottom) "
    "feature importance bar chart for M_anthro. Basin boundaries are shown in an inset map.",
    "A 3\u00d73 panel figure. Three columns for three basins. Row 1: TWS time series with "
    "observed (markers) and reconstructed (line). Row 2: TWSC comparison \u2014 observed vs. "
    "M_nat prediction vs. M_anthro prediction. Row 3: horizontal bar chart of feature "
    "importance (5 features) for each basin.",
)

# --- 4.2 Climate-Driven ---
doc.add_heading("4.2. Climate-Driven TWS Variability", level=2)

add_paragraph(
    "Not all TWS declines are attributable to direct human water extraction. Our analysis "
    "identifies a substantial subset of basins where precipitation deficits and/or enhanced "
    "evapotranspiration are the primary drivers of TWS loss. In these basins, the \u0394R\u00b2 is "
    "near zero or slightly negative, indicating that the addition of anthropogenic variables "
    "provides no improvement in predictive skill."
)

add_paragraph(
    "For example, [basin names in drought-prone regions] show significant TWS decline "
    "trends driven predominantly by multi-year precipitation deficits associated with "
    "large-scale climate modes such as the El Ni\u00f1o\u2013Southern Oscillation (ENSO) and the "
    "Pacific Decadal Oscillation (PDO). In [other basins], increasing evapotranspiration "
    "driven by rising air temperatures appears to be the dominant mechanism, consistent "
    "with projections of intensified atmospheric water demand under continued warming "
    "(Jung et al., 2010; Zhang et al., 2016)."
)

# --- 4.3 Uncertainties ---
doc.add_heading("4.3. Uncertainties and Limitations", level=2)

add_paragraph(
    "Several sources of uncertainty affect the results of this study. First, the spatial "
    "resolution of GRACE (~300 km effective resolution) limits the ability to resolve "
    "sub-basin heterogeneity in TWS changes, particularly in smaller basins where signal "
    "leakage from neighboring regions may influence basin-average estimates. Second, the "
    "hydroclimate predictor datasets (ERA5 reanalysis, GLEAM ET) are themselves model "
    "products with inherent biases and uncertainties, which propagate into the Random "
    "Forest predictions. Third, the PCR-GLOBWB abstraction estimates are based on national "
    "and sub-national water use statistics that may be outdated or incomplete, particularly "
    "in data-sparse regions."
)

add_paragraph(
    "The Random Forest framework, while powerful in capturing nonlinear relationships, is "
    "not a physically constrained model and may conflate correlated drivers. The OOB "
    "permutation importance metric assumes feature independence, which is violated when "
    "predictors are correlated (e.g., P and Q). Future work should explore conditional "
    "permutation importance (Strobl et al., 2008) or SHAP (SHapley Additive exPlanations) "
    "values to provide more robust and interpretable feature attribution."
)

add_paragraph(
    "Finally, the 213-month analysis period (2002\u20132019), while representing the longest "
    "available satellite gravimetry record, may be insufficient to fully characterize "
    "decadal-scale climate variability modes, potentially aliasing multi-decadal signals "
    "as secular trends."
)

doc.add_page_break()

# ===========================================================================
# 5. CONCLUSIONS
# ===========================================================================
doc.add_heading("5. Conclusions", level=1)

add_paragraph(
    "This study presents a comprehensive, data-driven framework for identifying and "
    "attributing global Terrestrial Water Storage (TWS) decline trends across the world\u2019s "
    "103 largest river basins. Our principal conclusions are:"
)

add_paragraph(
    "(1) Widespread TWS decline: [XX] out of 103 basins ([XX]%) exhibit statistically "
    "significant negative TWS trends over the GRACE/GRACE-FO era (2002\u20132019), with a "
    "mean decline rate of [XX] cm/year among these basins.",
    space_after=3,
)

add_paragraph(
    "(2) Effective gap reconstruction: Random Forest-based gap-filling successfully "
    "reconstructs the GRACE\u2013GRACE-FO observational gap with a mean Out-of-Bag R\u00b2 of [XX] "
    "across all basins, enabling continuous trend analysis.",
    space_after=3,
)

add_paragraph(
    "(3) Anthropogenic attribution: The Twin Random Forest Attribution Framework "
    "demonstrates that groundwater and surface water abstractions provide significant "
    "additional explanatory power (\u0394R\u00b2 > [XX]) in [XX] basins, predominantly in South "
    "Asia, the Middle East, and [other regions], identifying these as hotspots of "
    "human-induced water depletion.",
    space_after=3,
)

add_paragraph(
    "(4) Climate-driven variability: In the remaining basins, TWS variability is "
    "predominantly explained by natural hydroclimate drivers\u2014principally precipitation "
    "deficits and enhanced evapotranspiration\u2014highlighting the combined impacts of "
    "climate variability and global warming on freshwater reserves.",
    space_after=3,
)

add_paragraph(
    "(5) Policy implications: The basin-specific attribution diagnostics produced by "
    "this framework can directly inform targeted water resource management strategies, "
    "distinguishing basins where demand-side interventions (regulation of groundwater "
    "pumping, improved irrigation efficiency) are most urgently needed from those where "
    "climate adaptation measures (drought preparedness, reservoir management) are the "
    "priority."
)

add_paragraph(
    "Future extensions of this work will incorporate SHAP-based feature attribution for "
    "improved interpretability, expand the analysis to sub-basin scales where data "
    "resolution permits, and extend the temporal coverage as the GRACE-FO record lengthens."
)

doc.add_page_break()

# ===========================================================================
# ACKNOWLEDGMENTS
# ===========================================================================
doc.add_heading("Acknowledgments", level=1)

add_paragraph(
    "[The authors acknowledge computational resources provided by the DIRAC Supercomputer "
    "at IISER Kolkata. GRACE/GRACE-FO data were obtained from [source]. ERA5 data were "
    "provided by ECMWF through the Copernicus Climate Data Store. GLEAM data were obtained "
    "from [source]. PCR-GLOBWB simulations were provided by [source]. This work was "
    "supported by [funding agency/grant number].]"
)

doc.add_page_break()

# ===========================================================================
# REFERENCES
# ===========================================================================
doc.add_heading("References", level=1)

references = [
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5\u201332. https://doi.org/10.1023/A:1010933404324",
    "Castle, S. L., Thomas, B. F., Reager, J. T., Rodell, M., Swenson, S. C., & Famiglietti, J. S. (2014). Groundwater depletion during drought threatens future water security of the Colorado River Basin. Geophysical Research Letters, 41(16), 5904\u20135911.",
    "D\u00f6ll, P., M\u00fcller Schmied, H., Schuh, C., Portmann, F. T., & Eicker, A. (2014). Global-scale assessment of groundwater depletion and related groundwater abstractions: Combining hydrological modeling with information from well observations and GRACE satellites. Water Resources Research, 50(7), 5698\u20135720.",
    "Famiglietti, J. S. (2014). The global groundwater crisis. Nature Climate Change, 4(11), 945\u2013948.",
    "Famiglietti, J. S., Lo, M., Ho, S. L., Bethune, J., Anderson, K. J., Syed, T. H., Swenson, S. C., de Linage, C. R., & Rodell, M. (2011). Satellites measure recent rates of groundwater depletion in California\u2019s Central Valley. Geophysical Research Letters, 38(3), L03403.",
    "Feng, W., Zhong, M., Lemoine, J. M., Biancale, R., Hsu, H. T., & Xia, J. (2013). Evaluation of groundwater depletion in North China using the Gravity Recovery and Climate Experiment (GRACE) data and ground-based measurements. Water Resources Research, 49(4), 2110\u20132118.",
    "Hamed, K. H., & Rao, A. R. (1998). A modified Mann-Kendall trend test for autocorrelated data. Journal of Hydrology, 204(1\u20134), 182\u2013196.",
    "Hersbach, H., et al. (2020). The ERA5 global reanalysis. Quarterly Journal of the Royal Meteorological Society, 146(730), 1999\u20132049.",
    "Humphrey, V., Gudmundsson, L., & Seneviratne, S. I. (2017). A global reconstruction of climate-driven subdecadal water storage variability. Geophysical Research Letters, 44(5), 2300\u20132309.",
    "Joodaki, G., Wahr, J., & Swenson, S. (2014). Estimating the human contribution to groundwater depletion in the Middle East, from GRACE data, land surface models, and well observations. Water Resources Research, 50(3), 2679\u20132692.",
    "Jung, M., et al. (2010). Recent decline in the global land evapotranspiration trend due to limited moisture supply. Nature, 467(7318), 951\u2013954.",
    "Landerer, F. W., & Swenson, S. C. (2012). Accuracy of scaled GRACE terrestrial water storage estimates. Water Resources Research, 48(4), W04531.",
    "Martens, B., et al. (2017). GLEAM v3: satellite-based land evaporation and root-zone soil moisture. Geoscientific Model Development, 10(5), 1903\u20131925.",
    "Miralles, D. G., Holmes, T. R. H., De Jeu, R. A. M., Gash, J. H., Meesters, A. G. C. A., & Dolman, A. J. (2011). Global land-surface evaporation estimated from satellite-based observations. Hydrology and Earth System Sciences, 15(2), 453\u2013469.",
    "Rodell, M., Velicogna, I., & Famiglietti, J. S. (2009). Satellite-based estimates of groundwater depletion in India. Nature, 460(7258), 999\u20131002.",
    "Rodell, M., et al. (2018). Emerging trends in global freshwater availability. Nature, 557(7707), 651\u2013659.",
    "Scanlon, B. R., et al. (2018). Global models underestimate large decadal declining and rising water storage trends relative to GRACE satellite data. Proceedings of the National Academy of Sciences, 115(6), E1080\u2013E1089.",
    "Sen, P. K. (1968). Estimates of the regression coefficient based on Kendall\u2019s tau. Journal of the American Statistical Association, 63(324), 1379\u20131389.",
    "Strobl, C., Boulesteix, A. L., Kneib, T., Augustin, T., & Zeileis, A. (2008). Conditional variable importance for random forests. BMC Bioinformatics, 9, 307.",
    "Sutanudjaja, E. H., et al. (2018). PCR-GLOBWB 2: a 5 arcmin global hydrological and water resources model. Geoscientific Model Development, 11(6), 2429\u20132453.",
    "Tapley, B. D., Bettadpur, S., Ries, J. C., Thompson, P. F., & Watkins, M. M. (2004). GRACE measurements of mass variability in the Earth system. Science, 305(5683), 503\u2013505.",
    "Theil, H. (1950). A rank-invariant method of linear and polynomial regression analysis. Proceedings of the Royal Netherlands Academy of Arts and Sciences, 53, 386\u2013392, 521\u2013525, 1397\u20131412.",
    "Tiwari, V. M., Wahr, J., & Swenson, S. (2009). Dwindling groundwater resources in northern India, from satellite gravity observations. Geophysical Research Letters, 36(18), L18401.",
    "Voss, K. A., Famiglietti, J. S., Lo, M., De Linage, C., Rodell, M., & Swenson, S. C. (2013). Groundwater depletion in the Middle East from GRACE with implications for transboundary water management in the Tigris-Euphrates-Western Iran region. Water Resources Research, 49(2), 904\u2013914.",
    "Zaitchik, B. F., Rodell, M., & Reichle, R. H. (2008). Assimilation of GRACE terrestrial water storage data into a land surface model: Results for the Mississippi River basin. Journal of Hydrometeorology, 9(3), 535\u2013548.",
    "Zhang, Y., et al. (2016). Multi-decadal trends in global terrestrial evapotranspiration and its components. Scientific Reports, 6, 19124.",
]

for ref in references:
    ref_p = doc.add_paragraph()
    ref_p.paragraph_format.left_indent = Cm(1.27)
    ref_p.paragraph_format.first_line_indent = Cm(-1.27)
    ref_p.paragraph_format.space_after = Pt(3)
    ref_run = ref_p.add_run(ref)
    ref_run.font.name = "Times New Roman"
    ref_run.font.size = Pt(11)

doc.add_page_break()

# ===========================================================================
# SUPPLEMENTARY INFORMATION (Brief outline)
# ===========================================================================
doc.add_heading("Supplementary Information", level=1)

add_paragraph(
    "The following supplementary materials accompany this manuscript:",
    space_after=3,
)

add_paragraph(
    "Table S1: Complete basin-by-basin summary table including Basin ID, TWS trend "
    "(cm/year), Modified Mann-Kendall p-value, significance flag, Block CV NSE (M_nat "
    "and M_anthro), Block CV KGE (M_nat and M_anthro), OOB R\u00b2 (M_nat and M_anthro), \u0394R\u00b2, "
    "and dominant driver category for all 103 basins.",
    space_after=3,
)

add_paragraph(
    "Figure S1: Individual basin time series of observed and reconstructed TWS anomalies "
    "for all 103 basins (multi-page panel figure).",
    space_after=3,
)

add_paragraph(
    "Figure S2: Global map of OOB R\u00b2 for the gap-filling Random Forest model, showing "
    "spatial variability in reconstruction quality.",
    space_after=3,
)

add_paragraph(
    "Figure S3: Correlation matrix of predictor variables (P, ET, Q, GW_abs, SW_abs) "
    "across representative basins, illustrating inter-predictor dependencies.",
    space_after=3,
)

add_paragraph(
    "Code and Data Availability: All MATLAB code used in this study is available at "
    "[repository URL]. Processed basin-level time series and attribution results are "
    "archived at [DOI/repository].",
)

# ===========================================================================
# SAVE DOCUMENT
# ===========================================================================
doc.save(OUTPUT_FILE)
print(f"[OK] Publication-quality DOCX generated: {OUTPUT_FILE}")
print(f"   File size: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")
